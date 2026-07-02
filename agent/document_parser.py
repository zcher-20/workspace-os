from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument


@dataclass
class Paragraph:
    text: str
    page: int


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    full_text: str
    paragraphs: list[Paragraph]
    page_count: int

    @property
    def word_count(self) -> int:
        return len(self.full_text.split())


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def parse_document(filepath: Path) -> ParsedDocument:
    ext = filepath.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(filepath)
    if ext == ".docx":
        return _parse_docx(filepath)
    if ext == ".txt":
        return _parse_txt(filepath)
    raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")


def _parse_pdf(filepath: Path) -> ParsedDocument:
    paragraphs: list[Paragraph] = []
    with pdfplumber.open(filepath) as pdf:
        page_count = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text:
                for chunk in text.split("\n\n"):
                    chunk = chunk.strip()
                    if chunk:
                        paragraphs.append(Paragraph(text=chunk, page=page_num))

    full_text = "\n\n".join(p.text for p in paragraphs)
    return ParsedDocument(
        filename=filepath.name,
        file_type="pdf",
        full_text=full_text,
        paragraphs=paragraphs,
        page_count=page_count,
    )


def _parse_docx(filepath: Path) -> ParsedDocument:
    doc = DocxDocument(str(filepath))
    paragraphs = [
        Paragraph(text=p.text.strip(), page=1)
        for p in doc.paragraphs if p.text.strip()
    ]
    full_text = "\n\n".join(p.text for p in paragraphs)
    return ParsedDocument(
        filename=filepath.name,
        file_type="docx",
        full_text=full_text,
        paragraphs=paragraphs,
        page_count=1,
    )


def _parse_txt(filepath: Path) -> ParsedDocument:
    text = filepath.read_text(encoding="utf-8", errors="replace")
    paragraphs = [
        Paragraph(text=chunk.strip(), page=1)
        for chunk in text.split("\n\n") if chunk.strip()
    ]
    return ParsedDocument(
        filename=filepath.name,
        file_type="txt",
        full_text=text,
        paragraphs=paragraphs,
        page_count=1,
    )
